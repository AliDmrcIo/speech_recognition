import os
import platform
import shutil
os.environ['TORCH_FORCE_NO_WEIGHTS_ONLY_LOAD'] = 'true'
if platform.system() == "Windows":
    # Local development
    ffmpeg_path = r'C:\ffmpeg\bin'
    if os.path.exists(ffmpeg_path):
        os.environ['PATH'] = ffmpeg_path + ';' + os.environ['PATH']
elif platform.system() == "Linux":
    # AWS
    if not shutil.which("ffmpeg"):
        print("ffmpeg bulunamadı! Lütfen kurun:  sudo apt-get install -y ffmpeg")

from faster_whisper import WhisperModel
import numpy as np
import librosa
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
from docx import Document
from docx.shared import RGBColor, Pt
import colorsys
import time
import torch
from omegaconf.listconfig import ListConfig
torch.serialization.add_safe_globals([ListConfig])
from io import BytesIO
import whisperx
import gc
from whisperx.diarize import DiarizationPipeline
from docx import Document
from docx.shared import RGBColor, Pt
import colorsys
import streamlit as st
from dotenv import load_dotenv

load_dotenv()

def slow_model(path):
    device = "cpu"
    audio_file = path
    batch_size = 4
    compute_type = "int8"

    # 1. Transcribe
    print("--- 1. Transkripsyon yapılıyor ---")
    model = load_whisperx_model(device, compute_type)
    audio = whisperx.load_audio(audio_file)
    result = model.transcribe(audio, batch_size=batch_size)

    # 2. Align
    print("--- 2. Hizalama yapılıyor ---")
    model_a, metadata = load_align_model("en", device)
    result = whisperx.align(result["segments"], model_a, metadata, audio, device, return_char_alignments=False)

    # 3. Diarization
    print("--- 3. Konuşmacı ayrıştırılıyor ---")
    diarize_model = load_diarization_model(device)
    diarize_segments = diarize_model(audio)
    result = whisperx.assign_word_speakers(diarize_segments, result)

    # 4. Konuşmacıları tespit et
    print("--- 4. Konuşmacılar tespit ediliyor ---")
    speakers = set()
    for segment in result["segments"]:
        if "speaker" in segment:
            speakers.add(segment["speaker"])

    speakers = sorted(list(speakers))
    num_speakers = len(speakers)
    print(f"{num_speakers} konuşmacı tespit edildi:  {speakers}")

    # 5. Otomatik renk oluştur
    def generate_distinct_colors(n):
        colors = []
        for i in range(n):
            hue = i / n
            saturation = 0.8
            value = 0.9
            r, g, b = colorsys.hsv_to_rgb(hue, saturation, value)
            colors.append({
                'rgb': RGBColor(int(r * 255), int(g * 255), int(b * 255)),
                'r': int(r * 255),
                'g': int(g * 255),
                'b': int(b * 255)
            })
        return colors

    speaker_colors_list = generate_distinct_colors(num_speakers)
    speaker_colors = {speaker:  color for speaker, color in zip(speakers, speaker_colors_list)}
    speaker_names = {speaker:  f"Speaker {chr(65 + i)}" for i, speaker in enumerate(speakers)}

    print("Renkler atandı:")
    for speaker, name in speaker_names.items():
        color_info = speaker_colors[speaker]
        print(f"  {name}: RGB({color_info['r']}, {color_info['g']}, {color_info['b']})")

    # 6. Word formatında kaydet
    print("\n--- 5. Word dosyası oluşturuluyor ---")
    doc = Document()
    doc.add_heading('Speech Transcript', 0)

    # Konuşmacı bilgisi ekle
    info_para = doc.add_paragraph()
    info_para.add_run(f"• Total Number of Speakers: {num_speakers}\n").bold = True
    info_para.add_run(f"• Speakers: {', '.join(speaker_names.values())}\n")
    if result["segments"]:
        total_duration = float(result['segments'][-1]['end'])
        info_para.add_run(f"• Total Duration:  {total_duration:.1f} Seconds\n")
    doc.add_paragraph("─" * 60)

    # Segmentleri grupla
    grouped_segments = []
    current_group = None

    for segment in result["segments"]: 
        if "speaker" not in segment:
            continue
        
        speaker = segment.get("speaker")
        text = segment.get("text", "").strip()
        start = segment.get("start", 0)
        end = segment.get("end", 0)
        
        if not text:
            continue
        
        if current_group is None or current_group["speaker"] != speaker: 
            if current_group: 
                grouped_segments.append(current_group)
            current_group = {
                "speaker": speaker,
                "start": start,
                "end": end,
                "text": text
            }
        else:
            current_group["end"] = end
            current_group["text"] += " " + text

    if current_group:
        grouped_segments.append(current_group)

    # Word'e yaz
    for group in grouped_segments:
        speaker = group["speaker"]
        speaker_name = speaker_names.get(speaker, speaker)
        start_time = float(group["start"])
        end_time = float(group["end"])
        text = group["text"]
        color_info = speaker_colors.get(speaker, {'rgb': RGBColor(0, 0, 0)})
        
        paragraph = doc.add_paragraph()
        
        # Konuşmacı adı
        run_speaker = paragraph.add_run(f"🗣️ {speaker_name} ")
        run_speaker.font.bold = True
        run_speaker.font.size = Pt(12)
        run_speaker.font.color.rgb = color_info['rgb']
        
        # Zaman damgası
        run_time = paragraph.add_run(f"[{start_time:.1f}s-{end_time:.1f}s]:  ")
        run_time.font.size = Pt(11)
        run_time.font.color.rgb = RGBColor(128, 128, 128)
        
        # Konuşma metni
        run_text = paragraph.add_run(text)
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = color_info['rgb']
        
        paragraph.add_run("\n")

    # Kaydet
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def fast_model(path):
    device = "cpu"
    audio_file = rf"{path}"
    compute_type = "int8"

    # ============== 1. ASR ==============
    start = time.time()

    model = load_faster_whisper_model(device, compute_type)
    segments_raw, info = model.transcribe(
        audio_file,
        beam_size=1,
        language="en",
        vad_filter=True,
        vad_parameters=dict(min_silence_duration_ms=500)
    )

    result = {"segments": []}
    for segment in segments_raw: 
        result["segments"].append({
            "start": segment.start,
            "end": segment.end,
            "text": segment.text
        })

    print(f"✅ ASR tamamlandı: {time.time() - start:.2f}s ({len(result['segments'])} segment)")

    # ============== 2.  FEATURE EXTRACTION ==============
    start = time.time()

    y, sr = librosa.load(audio_file, sr=16000, mono=True)

    features = []
    for segment in result["segments"]:
        start_sample = int(segment["start"] * sr)
        end_sample = int(segment["end"] * sr)
        
        if start_sample >= len(y) or end_sample > len(y):
            end_sample = len(y)
        
        chunk = y[start_sample:end_sample]
        
        if len(chunk) > 512:
            # Pitch
            pitches, magnitudes = librosa.piptrack(y=chunk, sr=sr)
            pitch_mean = np.mean(pitches[pitches > 0]) if np.any(pitches > 0) else 0
            
            # MFCC
            mfcc = librosa.feature.mfcc(y=chunk, sr=sr, n_mfcc=13)
            mfcc_mean = np.mean(mfcc, axis=1)
            
            # Spectral features
            spectral_centroid = np.mean(librosa.feature.spectral_centroid(y=chunk, sr=sr))
            zcr = np.mean(librosa.feature.zero_crossing_rate(chunk))
            rms = np.mean(librosa.feature.rms(y=chunk))
            
            feature_vector = np.concatenate([
                [pitch_mean, spectral_centroid, zcr, rms],
                mfcc_mean
            ])
            
            features.append(feature_vector)
        else:
            if features:
                features.append(features[-1])
            else:
                features.append(np.zeros(17))

    features = np.array(features)

    # ============== 3. OTOMATİK KONUŞMACI SAYISI TESPİTİ ==============

    # Silhouette skoru ile optimal cluster sayısını bul
    best_n_speakers = 2  # varsayılan
    best_score = -1

    # 2-5 arası konuşmacı sayısını dene
    for n in range(2, min(6, len(features))):
        if len(features) > n:
            kmeans_test = KMeans(n_clusters=n, random_state=42, n_init=10)
            labels_test = kmeans_test.fit_predict(features)
            
            # Silhouette score hesapla (ne kadar yüksekse o kadar iyi)
            score = silhouette_score(features, labels_test)
            
            print(f"{n} konuşmacı → silhouette score: {score:.3f}")
            
            if score > best_score:
                best_score = score
                best_n_speakers = n

    NUM_SPEAKERS = best_n_speakers
    print(f"   ✅ Otomatik tespit: {NUM_SPEAKERS} konuşmacı (score: {best_score:.3f})")

    # ============== 4. CLUSTERING ==============
    print(f"\n--- 3. Diarization ({NUM_SPEAKERS} konuşmacı) ---")

    if len(features) > NUM_SPEAKERS:
        kmeans = KMeans(n_clusters=NUM_SPEAKERS, random_state=42, n_init=10)
        speaker_labels = kmeans.fit_predict(features)
        
        for i, segment in enumerate(result["segments"]):
            segment["speaker"] = f"SPEAKER_{speaker_labels[i]: 02d}"
    else:
        for i, segment in enumerate(result["segments"]):
            segment["speaker"] = f"SPEAKER_{i % NUM_SPEAKERS:02d}"

    # ============== 5. POST-PROCESSING (Smoothing) ==============
    print("   Post-processing (smoothing)...")

    # Çok kısa segmentleri önceki/sonraki speaker'a ata
    for i in range(1, len(result["segments"]) - 1):
        current = result["segments"][i]
        prev_seg = result["segments"][i-1]
        next_seg = result["segments"][i+1]
        
        duration = current["end"] - current["start"]
        
        # Çok kısa segment (<2 saniye) ve komşuları aynı speaker ise
        if duration < 2.0 and prev_seg["speaker"] == next_seg["speaker"]: 
            current["speaker"] = prev_seg["speaker"]

    # Majority voting
    smoothed_segments = []
    from collections import Counter

    for i in range(len(result["segments"])):
        current = result["segments"][i]
        
        # 5'li pencerede majority vote
        if i >= 2 and i < len(result["segments"]) - 2:
            window = [
                result["segments"][i-2]["speaker"],
                result["segments"][i-1]["speaker"],
                current["speaker"],
                result["segments"][i+1]["speaker"],
                result["segments"][i+2]["speaker"]
            ]
            most_common = Counter(window).most_common(1)[0][0]
            
            if window.count(current["speaker"]) == 1:
                current["speaker"] = most_common
        
        smoothed_segments.append(current)

    result["segments"] = smoothed_segments

    print(f"✅ Diarization tamamlandı: {time.time() - start:.2f}s")

    # ============== 6. KONUŞMACILAR ==============
    speakers = sorted(set(seg["speaker"] for seg in result["segments"]))
    num_speakers = len(speakers)
    print(f"✅ {num_speakers} konuşmacı tespit edildi:  {speakers}")

    # ============== 7. RENKLER ==============
    def generate_distinct_colors(n):
        colors = []
        for i in range(n):
            hue = i / n
            r, g, b = colorsys.hsv_to_rgb(hue, 0.8, 0.9)
            colors.append({
                'rgb': RGBColor(int(r * 255), int(g * 255), int(b * 255)),
                'r': int(r * 255),
                'g': int(g * 255),
                'b': int(b * 255)
            })
        return colors

    speaker_colors = {spk: color for spk, color in zip(speakers, generate_distinct_colors(num_speakers))}
    speaker_names = {spk: f"Speaker {chr(65 + i)}" for i, spk in enumerate(speakers)}

    print("\n🎨 Renkler atandı:")
    for speaker, name in speaker_names.items():
        color_info = speaker_colors[speaker]
        print(f"  {name}: RGB({color_info['r']}, {color_info['g']}, {color_info['b']})")

    # ============== 8. WORD DOSYASI ==============
    print("\n--- 4. Word dosyası oluşturuluyor ---")

    doc = Document()
    doc.add_heading('Speech Transcript', 0)

    info_para = doc.add_paragraph()
    info_para.add_run(f"• Total Speakers: {num_speakers}\n").bold = True
    info_para.add_run(f"• Speakers: {', '.join(speaker_names.values())}\n")
    if result["segments"]:
        total_duration = result['segments'][-1]['end']
        info_para.add_run(f"• Duration: {total_duration:.1f}s\n")
    doc.add_paragraph("─" * 60)

    # Segmentleri grupla
    grouped_segments = []
    current_group = None

    for seg in result["segments"]:
        if current_group is None or current_group["speaker"] != seg["speaker"]: 
            if current_group: 
                grouped_segments.append(current_group)
            current_group = {
                "speaker": seg["speaker"],
                "start":  seg["start"],
                "end": seg["end"],
                "text": seg["text"]
            }
        else:
            current_group["end"] = seg["end"]
            current_group["text"] += " " + seg["text"]

    if current_group: 
        grouped_segments.append(current_group)

    # Word'e yaz
    for group in grouped_segments:
        speaker = group["speaker"]
        speaker_name = speaker_names[speaker]
        color_info = speaker_colors[speaker]
        
        paragraph = doc.add_paragraph()
        
        run_speaker = paragraph.add_run(f"🗣️ {speaker_name} ")
        run_speaker.font.bold = True
        run_speaker.font.size = Pt(12)
        run_speaker.font.color.rgb = color_info['rgb']
        
        run_time = paragraph.add_run(f"[{group['start']:.1f}s-{group['end']:.1f}s]:  ")
        run_time.font.size = Pt(11)
        run_time.font.color.rgb = RGBColor(128, 128, 128)
        
        run_text = paragraph.add_run(group["text"])
        run_text.font.size = Pt(11)
        run_text.font.color.rgb = color_info['rgb']

    # Kaydet
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

@st.cache_resource
def load_whisperx_model(device="cpu", compute_type="int8"):
    import whisperx
    print("🔄 WhisperX modeli yükleniyor...")
    model = whisperx.load_model("base", device, compute_type=compute_type)
    print("✅ WhisperX modeli cache'lendi!")
    return model

@st.cache_resource
def load_faster_whisper_model(device="cpu", compute_type="int8"):
    from faster_whisper import WhisperModel
    print("Faster-Whisper modeli yükleniyor...")
    model = WhisperModel("tiny.en", device=device, compute_type=compute_type, cpu_threads=4)
    print("Faster-Whisper modeli cache'lendi!")
    return model

@st.cache_resource
def load_diarization_model(device="cpu"):
    from whisperx.diarize import DiarizationPipeline
    print("Diarization modeli yükleniyor...")
    model = DiarizationPipeline(
        use_auth_token=os.getenv("HUGGING_FACE_TOKEN"), 
        device=device
    )
    print("Diarization modeli cache'lendi!")
    return model

@st.cache_resource
def load_align_model(language_code, device="cpu"):
    import whisperx
    print(f"Align modeli yükleniyor ({language_code})...")
    model, metadata = whisperx.load_align_model(language_code=language_code, device=device)
    print("Align modeli cache'lendi!")
    return model, metadata
